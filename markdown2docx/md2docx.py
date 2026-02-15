# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.shared import Inches
import text_template as template
import argparse
import os
import datetime

def log(message, level="INFO"):
    """打印带时间戳的日志信息"""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] [{level}] {message}")

# 全局变量
debug_mode = False

def debug(message):
    """打印调试信息（仅在debug_mode为True时）"""
    if debug_mode:
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"[{timestamp}] [DEBUG] {message}")

def match_markdown_images(text):
    # 正则表达式用于匹配 Markdown 中的![]() 图像标记
    pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    matches = pattern.findall(text)
    results = []
    for match in matches:
        # match 是一个元组，第一个元素是中括号内的内容，第二个元素是括号内的内容
        results.append((match[0], match[1]))
    return results

def match_markdown_links(text):
    # 正则表达式用于匹配 Markdown 中的 []() 链接
    pattern = re.compile(r'\[(.*?)\]\((.*?)\)')
    matches = pattern.findall(text)
    results = []
    for match in matches:
        # match 是一个元组，第一个元素是中括号内的内容，第二个元素是括号内的内容
        results.append((match[0], match[1]))
    return results
def is_url(url):
    """验证字符串是否为有效的URL"""
    if not url or not isinstance(url, str):
        return False
    
    # 增强的 URL 正则表达式
    pattern = re.compile(
        r'^(https?|ftp)://'  # 协议部分，支持 http, https, ftp
        r'([A-Za-z0-9][A-Za-z0-9-]{0,61}[A-Za-z0-9]\.)+'  # 域名部分，支持多级域名
        r'([A-Za-z]{2,6}|[A-Za-z0-9-]{2,})'  # 顶级域名
        r'(:[0-9]{1,5})?'  # 端口部分，可选，1-5位数字
        r'(/[^\s]*)?'  # 路径部分，可选，不包含空白字符
        r'(\?[^\s]*)?'  # 查询部分，可选，不包含空白字符
        r'(#[^\s]*)?$'  # 片段部分，可选，不包含空白字符
    )
    
    # 尝试匹配
    if pattern.match(url):
        return True
    
    # 额外检查：确保URL长度合理
    if len(url) > 2048:
        return False
    
    return False

def write_content(text):
    document.add_paragraph(text)

def write_pic(path):
    debug(f"写入图片路径 ===> {path} ")
    try:
        # 检查图片文件是否存在
        if not os.path.exists(path):
            log(f"图片文件不存在 '{path}'，跳过添加图片", "WARNING")
            return
        
        # 检查文件是否为图片格式
        img_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
        if not any(path.lower().endswith(ext) for ext in img_extensions):
            log(f"文件可能不是图片格式 '{path}'，尝试添加", "WARNING")
        
        document.add_picture(path)
        debug(f"图片添加成功：{path}")
    except FileNotFoundError as e:
        log(f"图片文件未找到 '{path}'：{e}", "ERROR")
    except PermissionError as e:
        log(f"没有权限访问图片文件 '{path}'：{e}", "ERROR")
    except Exception as e:
        log(f"添加图片时发生异常 '{path}'：{e}", "ERROR")

def write_end():
    try:
        document.add_page_break()
        
        # 检查输出目录是否存在
        output_path = output_name + '.docx'
        output_dir = os.path.dirname(output_path)
        
        # 如果指定了目录且目录不存在，尝试创建
        if output_dir and not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
                log(f"创建输出目录：{output_dir}")
            except Exception as e:
                log(f"无法创建输出目录 '{output_dir}'：{e}", "WARNING")
        
        document.save(output_path)
        log(f"文档保存成功：{output_path}", "SUCCESS")
    except PermissionError as e:
        log(f"没有权限保存文件 '{output_name}.docx'：{e}", "ERROR")
    except Exception as e:
        log(f"保存文档时发生异常：{e}", "ERROR")

def read_markdown_file(file_path,write,writeEnd):
    try:
        log(f"开始读取Markdown文件：{file_path}")
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            log(f"文件不存在 '{file_path}'", "ERROR")
            return
        
        # 检查文件是否可读
        if not os.access(file_path, os.R_OK):
            log(f"没有权限读取文件 '{file_path}'", "ERROR")
            return
        
        with open(file_path, 'r', encoding='utf-8') as file:
            isArtiles = False
            artContents = ''
            line_count = 0
            
            try:
                for line_number, line in enumerate(file, start=1):
                    line_count += 1
                    # 去除行尾的换行符
                    line = line.strip().rstrip('\n')
                    if len(line) > 0:
                        ## 标题部分
                        if(line.startswith('## ')):
                            try:
                                title = line.replace('## ', "").strip()
                                # 是文章标题情况下
                                if(title == '📕 精选文章'):
                                    isArtiles = True
                                else: 
                                    if isArtiles: 
                                        try:
                                            content = template.templateRedUl.substitute(ul = artContents)
                                            write(content.strip())
                                        except Exception as e:
                                            log(f"处理文章内容时发生异常（行 {line_number}）：{e}", "ERROR")
                                    isArtiles = False
                                try:
                                    newline = template.templateTitleH2.substitute(h2 = title)
                                    ## 写主题TXT
                                    write(newline.strip())
                                except Exception as e:
                                    log(f"处理标题模板时发生异常（行 {line_number}）：{e}", "ERROR")
                            except Exception as e:
                                log(f"处理标题时发生异常（行 {line_number}）：{e}", "ERROR")
                        ## 内容部分
                        else:
                            ## 文章主题内容
                            if isArtiles: 
                                try:
                                    line = line.replace('* ', "").strip()
                                    results = match_markdown_links(line.strip())
                                    if results:
                                        newline = template.templateArtileLi.substitute({'li': '📄' + results[0][0],'url':'🔗【' +results[0][1]+'】'})
                                        artContents += newline.strip()
                                    else:
                                        log(f"未找到链接（行 {line_number}）：{line}", "WARNING")
                                        artContents += line
                                except Exception as e:
                                    log(f"处理文章内容时发生异常（行 {line_number}）：{e}", "ERROR")
                            else:
                            ## 其他主题内容
                            
                                if(line.startswith('**')):
                                    try:
                                        title = line.replace('*', "").strip()
                                        title = template.templateTitleBgH3.substitute(h3 = '# ' + title)
                                        write(title.strip())
                                    except Exception as e:
                                        log(f"处理加粗标题时发生异常（行 {line_number}）：{e}", "ERROR")
                                elif (line.strip().startswith('* [')):
                                    try:
                                        line = line.replace('* ', "").strip()
                                        results = match_markdown_links(line)
                                        if results:
                                            write(results[0][0] + ":" + results[0][1])    
                                        else:
                                            log(f"未找到链接（行 {line_number}）：{line}", "WARNING")
                                    except Exception as e:
                                        log(f"处理列表链接时发生异常（行 {line_number}）：{e}", "ERROR")
                                elif is_url(line):
                                    try:
                                        url = template.templateNetUrl.substitute(url = '🔗【' + line + '】')
                                        write(url.strip())
                                    except Exception as e:
                                        log(f"处理URL时发生异常（行 {line_number}）：{e}", "ERROR")
                                elif (line.strip().startswith('![')):
                                    try:
                                        imgPath = match_markdown_images(line)
                                        if imgPath:
                                            img_full_path = file_path.rsplit('/', 1)[0]+ "/" + imgPath[0][1]
                                            write_pic(img_full_path)
                                        else:
                                            log(f"未找到图片路径（行 {line_number}）：{line}", "WARNING")
                                    except Exception as e:
                                        log(f"处理图片时发生异常（行 {line_number}）：{e}", "ERROR")
                                else:
                                    try:
                                        newLine = template.templateContent.substitute(content = line.strip())
                                        write(newLine.strip())
                                    except Exception as e:
                                        log(f"处理内容模板时发生异常（行 {line_number}）：{e}", "ERROR")
            except UnicodeDecodeError as e:
                log(f"文件编码错误 '{file_path}'：{e}", "ERROR")
            except Exception as e:
                log(f"读取文件内容时发生异常（行 {line_number}）：{e}", "ERROR")
            
            # 添加结尾内容
            try:
                write('<br/>')
                write('<br/>')
                ## 宣传语 
                write(template.templateDividerLine) 
                write(template.templateMiddleTitle.substitute(content = "你的关注是我更新的最大动力😙\n💪🏻基本每周更新~").strip())
                write(template.templateDividerLine) 
                ## 公众号的二维码
                write(template.templateWXCard)
                write_pic("../docs/wchat/julystudio.jpg")
                writeEnd()
            except Exception as e:
                log(f"添加结尾内容时发生异常：{e}", "ERROR")
        
        log(f"文件读取完成，共处理 {line_count} 行")
        
    except FileNotFoundError as e:
        log(f"文件未找到 '{file_path}'：{e}", "ERROR")
    except PermissionError as e:
        log(f"没有权限访问文件 '{file_path}'：{e}", "ERROR")
    except UnicodeDecodeError as e:
        log(f"文件编码错误 '{file_path}'：{e}", "ERROR")
    except Exception as e:
        log(f"读取Markdown文件时发生异常：{e}", "ERROR")

def md2docx(file_path):
    """将Markdown文件转换为docx文件"""
    try:
        log(f"开始转换Markdown文件：{file_path}")
        
        # 验证文件路径
        if not file_path:
            log("文件路径不能为空", "ERROR")
            return False
        
        if not os.path.exists(file_path):
            log(f"文件不存在 '{file_path}'", "ERROR")
            return False
        
        if not os.path.isfile(file_path):
            log(f"路径不是文件 '{file_path}'", "ERROR")
            return False
        
        # 创建文档对象
        document = Document()
        log("文档对象创建成功")
        
        # 这里可以添加更多初始化操作
        
        return True
        
    except ImportError as e:
        log(f"缺少必要的库：{e}", "ERROR")
        return False
    except Exception as e:
        log(f"初始化转换时发生异常：{e}", "ERROR")
        return False

if __name__ == "__main__":
    # 启用调试模式
    # global debug_mode
    
    # 在这里输入你要读取的文件的路径
    parser = argparse.ArgumentParser(description='Markdown转docx工具')
    parser.add_argument('-input_path',type=str, required=True, help='Markdown文件路径')
    parser.add_argument('-output_name',type=str, required=True, help='输出docx文件名')
    parser.add_argument('-debug', action='store_true', help='启用调试模式')
    args = parser.parse_args()
    
    # 启用调试模式
    debug_mode = args.debug
    
    try:
        file_path = args.input_path
        output_name = args.output_name
        
        log(f"开始执行Markdown转docx工具")
        log(f"输入文件路径：{file_path}")
        log(f"输出文件名：{output_name}")
        
        # 验证输入文件是否存在
        if not os.path.exists(file_path):
            log(f"输入文件 '{file_path}' 不存在！", "ERROR")
            exit(1)
        
        # 验证输入文件是否为Markdown文件
        if not file_path.endswith('.md'):
            log(f"输入文件 '{file_path}' 可能不是Markdown文件！", "WARNING")
        
        document = Document()
        read_markdown_file(file_path,write_content,write_end)
        log(f"转换完成：{output_name}.docx", "SUCCESS")
    except Exception as e:
        log(f"命令行参数处理错误：{e}", "ERROR")
        exit(1)
    #python3 md2docx.py -input_path ../Weekly/No21/No21.md -output_name 余小余周刊-第21期